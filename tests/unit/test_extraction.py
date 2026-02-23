"""Tests for the text extraction service."""

from __future__ import annotations

import pytest

from maia_vectordb.core.exceptions import ValidationError
from maia_vectordb.services.extraction import (
    detect_file_type,
    extract_text,
    is_binary_format,
)


class TestDetectFileType:
    """Tests for detect_file_type()."""

    @pytest.mark.parametrize(
        "filename,expected",
        [
            ("doc.pdf", ".pdf"),
            ("report.DOCX", ".docx"),
            ("notes.txt", ".txt"),
            ("data.JSON", ".json"),
            ("page.HTML", ".html"),
            ("style.csv", ".csv"),
            ("feed.xml", ".xml"),
            ("config.yaml", ".yaml"),
            ("config.yml", ".yml"),
            ("readme.md", ".md"),
            ("page.htm", ".htm"),
        ],
    )
    def test_supported_extensions(self, filename: str, expected: str) -> None:
        assert detect_file_type(filename) == expected

    def test_unsupported_extension_raises(self) -> None:
        with pytest.raises(ValidationError, match="Unsupported file format"):
            detect_file_type("image.png")

    def test_no_extension_defaults_to_txt(self) -> None:
        assert detect_file_type("README") == ".txt"

    def test_case_insensitive(self) -> None:
        assert detect_file_type("DOC.PDF") == ".pdf"


class TestIsBinaryFormat:
    def test_pdf_is_binary(self) -> None:
        assert is_binary_format(".pdf") is True

    def test_docx_is_binary(self) -> None:
        assert is_binary_format(".docx") is True

    def test_txt_is_not_binary(self) -> None:
        assert is_binary_format(".txt") is False


class TestExtractPdf:
    def test_valid_pdf(self) -> None:
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello from PDF")
        raw = doc.tobytes()
        doc.close()

        result = extract_text(raw, ".pdf")
        assert "Hello from PDF" in result

    def test_empty_pdf_raises(self) -> None:
        import fitz

        doc = fitz.open()
        doc.new_page()  # blank page with no text
        raw = doc.tobytes()
        doc.close()

        with pytest.raises(ValidationError, match="no extractable text"):
            extract_text(raw, ".pdf")

    def test_corrupt_pdf_raises_validation_error(self) -> None:
        """Corrupt/invalid PDF bytes should raise ValidationError, not 500."""
        with pytest.raises(ValidationError, match="Failed to parse PDF"):
            extract_text(b"this is not a pdf at all", ".pdf")


class TestExtractDocx:
    def test_valid_docx(self) -> None:
        import io

        import docx

        doc = docx.Document()
        doc.add_paragraph("Hello from DOCX")
        buf = io.BytesIO()
        doc.save(buf)
        raw = buf.getvalue()

        result = extract_text(raw, ".docx")
        assert "Hello from DOCX" in result

    def test_empty_docx_raises(self) -> None:
        import io

        import docx

        doc = docx.Document()
        buf = io.BytesIO()
        doc.save(buf)
        raw = buf.getvalue()

        with pytest.raises(ValidationError, match="no extractable text"):
            extract_text(raw, ".docx")

    def test_corrupt_docx_raises_validation_error(self) -> None:
        """Corrupt/invalid DOCX bytes should raise ValidationError, not 500."""
        with pytest.raises(ValidationError, match="Failed to parse DOCX"):
            extract_text(b"this is not a docx at all", ".docx")
